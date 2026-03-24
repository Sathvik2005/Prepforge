from pathlib import Path
import re

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt


def add_inline(paragraph, node):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    if node.name in {"strong", "b", "em", "i", "code", "a", "span"}:
        for child in node.children:
            run = paragraph.add_run("")
            if node.name in {"strong", "b"}:
                run.bold = True
            if node.name in {"em", "i"}:
                run.italic = True
            if node.name == "code":
                run.font.name = "Consolas"
            if node.name == "a":
                href = node.get("href", "")
                txt = node.get_text(strip=False)
                run.text = f"{txt} ({href})" if href else txt
                return

            if isinstance(child, NavigableString):
                run.text = str(child)
            else:
                # recurse nested formatting into the same paragraph
                add_inline(paragraph, child)
        return

    for child in node.children:
        add_inline(paragraph, child)


def add_paragraph_from_tag(doc: Document, tag: Tag):
    p = doc.add_paragraph()
    for child in tag.children:
        add_inline(p, child)


def add_list(doc: Document, tag: Tag, ordered: bool):
    style = "List Number" if ordered else "List Bullet"
    for li in tag.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        for child in li.contents:
            add_inline(p, child)


def add_table(doc: Document, tag: Tag):
    rows = tag.find_all("tr")
    if not rows:
        return

    first_cells = rows[0].find_all(["th", "td"])
    cols = len(first_cells)
    table = doc.add_table(rows=0, cols=max(cols, 1))
    table.style = "Table Grid"

    for tr in rows:
        cells = tr.find_all(["th", "td"])
        row = table.add_row().cells
        for i, cell in enumerate(cells[: len(row)]):
            row[i].text = re.sub(r"\s+", " ", cell.get_text(" ", strip=True))


def add_code_block(doc: Document, tag: Tag):
    code_text = tag.get_text("", strip=False)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def markdown_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    html = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "toc", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for elem in soup.contents:
        if isinstance(elem, NavigableString):
            continue

        name = elem.name.lower()
        if name == "h1":
            doc.add_heading(elem.get_text(" ", strip=True), level=1)
        elif name == "h2":
            doc.add_heading(elem.get_text(" ", strip=True), level=2)
        elif name == "h3":
            doc.add_heading(elem.get_text(" ", strip=True), level=3)
        elif name == "h4":
            doc.add_heading(elem.get_text(" ", strip=True), level=4)
        elif name == "p":
            add_paragraph_from_tag(doc, elem)
        elif name == "ul":
            add_list(doc, elem, ordered=False)
        elif name == "ol":
            add_list(doc, elem, ordered=True)
        elif name == "table":
            add_table(doc, elem)
        elif name == "pre":
            add_code_block(doc, elem)
        elif name == "hr":
            doc.add_paragraph("-" * 40)
        elif name == "blockquote":
            p = doc.add_paragraph(elem.get_text(" ", strip=True))
            p.style = "Intense Quote"
        else:
            # fallback: preserve text if tag is unhandled
            txt = elem.get_text(" ", strip=True)
            if txt:
                doc.add_paragraph(txt)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


if __name__ == "__main__":
    source = Path(r"e:\HackAura\prepwiser\docs\PREPWISER_PROJECT_REPORT_VIT_AP.md")
    target = Path(r"e:\HackAura\prepwiser\docs\PREPWISER_PROJECT_REPORT_VIT_AP.docx")
    markdown_to_docx(source, target)
    print(f"Created: {target}")
